from flask import request, jsonify, current_app, send_file, Blueprint, session, url_for
from flask_login import login_required, current_user
import os
import io
import json
import traceback
import tempfile  # Add this import
import fitz  # PyMuPDF
import re  # Added for regex pattern matching
from docx import Document
from docxtpl import DocxTemplate, RichText
from backend.services.case_service import get_case_by_id
from backend.schemas import case_schema  # Add this import
from backend.app.discovery import (
    parse_requests_for_production,
    parse_form_interrogatories,
    parse_special_interrogatories,
    parse_requests_for_admission,
    build_rfp_prompt,
    build_form_interrogatories_prompt,
    build_special_interrogatories_prompt,
    build_rfa_prompt,
    DiscoveryResponseService
)
from backend.models import Case
from backend.extensions import csrf
from . import bp
from backend.app.discovery.formatters import format_form_interrogatory_responses, format_medical_records
import logging
from datetime import datetime
from typing import Dict, Any
# Add these imports to the top of backend/api/discovery.py

# Modify this existing import line:
from backend.services.case_service import get_case_by_id

# To include CaseNotFoundError:
from backend.services.case_service import get_case_by_id, CaseNotFoundError

# Add this new import:
from backend.app.discovery.registry import get_discovery_type_info
from docx import Document as DocxDocument

# Function to strip markdown formatting
def strip_markdown(text):
    """Remove markdown formatting from text."""
    if not text:
        return ""
    # Remove bold markdown
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    # Remove italic markdown
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    # Remove headers (e.g., # Header)
    text = re.sub(r'^#+\s+', '', text, flags=re.MULTILINE)
    return text

# Test endpoint to verify the blueprint is working
@bp.route('/test', methods=['GET'])
def test_endpoint():
    """Simple test endpoint to verify the discovery blueprint is working."""
    print("DEBUG: Test endpoint called successfully")
    return jsonify({"message": "Discovery API is working", "authenticated": current_user.is_authenticated}), 200

def get_cached_objection_master():
    """Load and cache the objection master sheet as plain text."""
    if not hasattr(current_app, '_objection_master_cache'):
        objection_path = os.path.join(current_app.root_path, 'templates', 'MASTER SHEET Discovery Objection.docx')
        try:
            doc = DocxDocument(objection_path)
            objection_master = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
            current_app._objection_master_cache = objection_master
        except Exception as doc_error:
            print(f"DEBUG: Error loading objection master sheet: {str(doc_error)}")
            current_app._objection_master_cache = "Error loading objection master sheet"
    return current_app._objection_master_cache

@bp.route('/discovery/cases/<int:case_id>/parse', methods=['POST'])
@login_required
def parse_discovery_document(case_id):
    """
    Accepts a file upload and discovery_type, parses the file, and returns just the questions.
    This is step 1 of the two-step process for discovery document generation.
    """
    print(f"DEBUG: Discovery parse endpoint accessed by user {current_user.id} for case {case_id}")
    print(f"DEBUG: Request form data: {request.form}")
    print(f"DEBUG: Request files: {request.files.keys() if request.files else 'None'}")
    
    # Check for 'document' instead of 'file' to match Documents API pattern
    document = request.files.get('document')
    discovery_type = request.form.get('discovery_type')
    
    print(f"DEBUG: Extracted discovery_type: {discovery_type}")
    print(f"DEBUG: Extracted document: {document.filename if document else 'None'}")
    
    if not document or not discovery_type:
        print("DEBUG: Missing document or discovery_type, returning 400")
        return jsonify({'error': 'Missing document or discovery_type'}), 400

    # Create a temporary file with a secure name using tempfile module
    temp_file = None
    temp_path = None
    try:
        # Create a named temporary file that closes when done
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
            temp_path = temp_file.name
            print(f"DEBUG: Created temporary file at {temp_path}")
            
            # Save uploaded file content to the temporary file
            document.save(temp_file)
            temp_file.flush()  # Ensure all data is written
            
        print(f"DEBUG: File saved successfully, size: {os.path.getsize(temp_path) if os.path.exists(temp_path) else 'file not found'}")
        
        # Load case details
        print(f"DEBUG: Attempting to get case {case_id} for user {current_user.id}")
        case = get_case_by_id(case_id, user_id=current_user.id)
        print(f"DEBUG: Successfully retrieved case {case_id}")
        
        case_details = {}
        if hasattr(case, 'to_dict'):
            print("DEBUG: Case has to_dict method, calling it")
            case_details = case.to_dict()
        else:
            print("DEBUG: Case doesn't have to_dict method, creating basic dict")
            # Create a basic dict with essential attributes
            case_details = {
                'id': case.id,
                'display_name': case.display_name if hasattr(case, 'display_name') else 'Unknown',
                # Add other essential fields
            }
        
        # Load objection master sheet (as plain text)
        objection_master = get_cached_objection_master()
        print(f"DEBUG: Successfully loaded objection master sheet, length: {len(objection_master)}")
        
        # Use the orchestrator service - just to parse, not for full response
        print(f"DEBUG: About to initialize DiscoveryResponseService")
        service = DiscoveryResponseService()
        print(f"DEBUG: Service initialized successfully")
        
        print(f"DEBUG: Calling service.respond with discovery_type={discovery_type}")
        
        # Add some pre-processing checks to validate the PDF
        try:
            # Quick text extraction test to verify PDF is readable
            with fitz.open(temp_path) as doc:
                if len(doc) > 0:
                    sample_text = doc[0].get_text()[:200]
                    print(f"DEBUG: Sample text from first page: {sample_text}")
                else:
                    print("DEBUG: PDF appears to have no pages")
        except Exception as pdf_check_error:
            print(f"DEBUG: Error in PDF pre-check: {pdf_check_error}")
            
        # Get the AI response but only return questions
        result = service.respond(discovery_type, temp_path, case_details, objection_master)
        print(f"DEBUG: Service.respond completed successfully")
        
        # Extract just the questions and clean them for display
        questions = result.get('questions', [])
        cleaned_questions = []
        
        for q in questions:
            question_number = q.get('number', '')
            question_text = q.get('text', '')
            
            # Clean the question text (remove any existing headers)
            question_text = strip_markdown(question_text)
            # Remove any existing headers
            question_text = re.sub(r'^REQUEST FOR PRODUCTION NO\.\s*\d+\s*:\s*', '', question_text, flags=re.IGNORECASE)
            
            cleaned_questions.append({
                'id': f"q_{question_number}",  # Unique ID for frontend
                'number': question_number,
                'text': question_text.strip()
            })
            
        # Store the full result in the session for the second step
        # We need to ensure this data is accessible in the generate endpoint
        session_key = f"discovery_result_{case_id}_{current_user.id}"
        current_app.config[session_key] = result
        
        return jsonify({
            'questions': cleaned_questions,
            'discovery_type': discovery_type,
            'message': 'Document parsed successfully',
            'session_key': session_key  # Include this for reference
        }), 200
        
    except Exception as e:
        error_trace = traceback.format_exc()
        print(f"DEBUG: Error in parse_discovery_document: {str(e)}")
        print(f"DEBUG: Error traceback: {error_trace}")
        current_app.logger.error(f"Error processing discovery for case {case_id}: {e}")
        current_app.logger.error(f"Traceback: {error_trace}")
        return jsonify({"error": f"An error occurred: {str(e)}"}), 500
    finally:
        # Clean up the temporary file
        if temp_file and temp_path and os.path.exists(temp_path):
            try:
                os.remove(temp_path)
                print(f"DEBUG: Temporary file {temp_path} removed")
            except Exception as cleanup_error:
                print(f"DEBUG: Error removing temporary file: {str(cleanup_error)}")


@bp.route('/discovery/cases/<int:case_id>/generate-document', methods=['POST'])
@login_required
def generate_discovery_document(case_id):
    """
    Unified endpoint to generate discovery documents for all types.
    Uses registry pattern to determine workflow and template.
    """
    print(f"DEBUG: Generate discovery document for case {case_id} by user {current_user.id}")
    
    try:
        # Get request data
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No JSON data provided'}), 400
            
        discovery_type = data.get('discovery_type')
        if not discovery_type:
            return jsonify({'error': 'discovery_type is required'}), 400
            
        print(f"DEBUG: Discovery type: {discovery_type}")
        
        # Get configuration from registry
        try:
            type_config = get_discovery_type_info(discovery_type)
        except ValueError as e:
            return jsonify({'error': str(e)}), 400
            
        print(f"DEBUG: Using template: {type_config['template_file']}")
        print(f"DEBUG: Workflow type: {type_config['workflow_type']}")
        
        # Verify case ownership
        case = get_case_by_id(case_id, user_id=current_user.id)
        
        # Handle different workflows based on registry
        if type_config['workflow_type'] == 'format_responses':
            # Form Interrogatories workflow
            return _generate_form_interrogatory_document(case, type_config, data)
            
        elif type_config['workflow_type'] == 'parse_and_select':
            # RFPs & Special Interrogatories workflow
            return _generate_parse_and_select_document(case, type_config, data)
            
        else:
            return jsonify({'error': f'Unknown workflow type: {type_config["workflow_type"]}'}), 500
            
    except CaseNotFoundError as e:
        return jsonify({'error': str(e)}), 404
    except Exception as e:
        print(f"DEBUG: Error generating discovery document: {str(e)}")
        traceback.print_exc()
        return jsonify({'error': f'Failed to generate document: {str(e)}'}), 500

def _generate_form_interrogatory_document(case, type_config, data):
    """Handle form interrogatory document generation."""
    print("DEBUG: Processing form interrogatory workflow")
    
    # Get session key from request
    session_key = data.get('session_key')
    if not session_key:
        return jsonify({'error': 'session_key is required'}), 400
        
    # Retrieve stored parsing result
    stored_result = current_app.config.get(session_key)
    if not stored_result:
        return jsonify({
            'error': 'Session data not found or expired. Please re-upload the document.'
        }), 400
        
    print(f"DEBUG: Found stored result with {len(stored_result.get('questions', []))} questions")
    
    # Get questions and AI response
    questions = stored_result.get('questions', [])
    ai_response = stored_result.get('ai_response', '')
    
    if not questions:
        return jsonify({'error': 'No questions found in stored result'}), 400
        
    # Process AI response
    responses_dict = _parse_ai_response_by_question(ai_response, type_config)
    
    # Prepare context for document generation
    context = _build_case_context(case)
    context['responses'] = responses_dict
    
    # Clean up session data
    if session_key in current_app.config:
        del current_app.config[session_key]
        print(f"DEBUG: Cleaned up session key: {session_key}")
    
    # Generate document
    return _render_and_send_document(
        template_name=type_config['template_file'],
        context=context,
        case_id=case.id,
        discovery_type='form_interrogatories'
    )

def _generate_parse_and_select_document(case, type_config, data):
    """Handle RFPs & Special Interrogatories workflow (parse → select → generate)."""
    print("DEBUG: Processing parse and select workflow")
    
    # Get session key and selections from request
    session_key = data.get('session_key')
    selections = data.get('selections', {})
    
    if not session_key:
        return jsonify({'error': 'session_key is required'}), 400
        
    # Retrieve stored parsing result
    stored_result = current_app.config.get(session_key)
    if not stored_result:
        return jsonify({
            'error': 'Session data not found or expired. Please re-upload the document.'
        }), 400
        
    print(f"DEBUG: Found stored result with {len(stored_result.get('questions', []))} questions")
    
    # Get questions and AI response
    questions = stored_result.get('questions', [])
    ai_response = stored_result.get('ai_response', '')
    
    if not questions:
        return jsonify({'error': 'No questions found in stored result'}), 400
        
    # Process AI response and apply user selections
    combined_responses = _process_ai_responses_with_selections(
        questions, ai_response, selections, type_config
    )
    
    # Prepare context for document generation
    context = _build_case_context(case)
    context['responses'] = combined_responses
    
    # Clean up session data
    if session_key in current_app.config:
        del current_app.config[session_key]
        print(f"DEBUG: Cleaned up session key: {session_key}")
    
    # Generate document
    return _render_and_send_document(
        template_name=type_config['template_file'],
        context=context,
        case_id=case.id,
        discovery_type=data.get('discovery_type', 'unknown')
    )

def _build_case_context(case):
    """Build standard case context for template rendering."""
    return {
        # Core case information
        'case_name': getattr(case, 'display_name', ''),
        'case_number': getattr(case, 'case_number', ''),
        'plaintiff': getattr(case, 'plaintiff', ''),
        'defendant': getattr(case, 'defendant', ''),
        'judge': getattr(case, 'judge', ''),
        'jurisdiction': getattr(case, 'jurisdiction', ''),
        'county': getattr(case, 'county', ''),
        
        # Date fields
        'filing_date': getattr(case, 'filing_date', ''),
        'trial_date': getattr(case, 'trial_date', ''),
        'incident_date': getattr(case, 'incident_date', ''),
        
        # Other fields
        'incident_location': getattr(case, 'incident_location', ''),
        'incident_description': getattr(case, 'incident_description', ''),
        'case_type': getattr(case, 'case_type', ''),
        'defendant_counsel_info': getattr(case, 'defendant_counsel_info', ''),
        'plaintiff_counsel_info': getattr(case, 'plaintiff_counsel_info', ''),
        'vehicle_details': getattr(case, 'vehicle_details', ''),
        'defendant_counsel_attorneys': getattr(case, 'defendant_counsel_attorneys', ''),
        'defendant_counsel_firm': getattr(case, 'defendant_counsel_firm', ''),
        'defendant_counsel_address': getattr(case, 'defendant_counsel_address', ''),
        'defendant_counsel_email': getattr(case, 'defendant_counsel_email', ''),
        'defendant_counsel_phone': getattr(case, 'defendant_counsel_phone', ''),
        'acting_attorney': getattr(case, 'acting_attorney', ''),
        'acting_clerk': getattr(case, 'acting_clerk', ''),
        
        # Current date and year
        'current_date': datetime.now().strftime('%B %d, %Y'),
        'current_year': datetime.now().year,
    }

def _process_ai_responses_with_selections(questions, ai_response, selections, type_config):
    """Process AI responses and combine with user selections."""
    from docxtpl import RichText
    
    # Different standard responses based on discovery type
    standard_responses = {
        'requests_for_production': {
            'will_provide': 'Plaintiff will produce responsive documents.',
            'none_found': 'Plaintiff has no responsive documents to produce.',
            'no_text': ''
        },
        'special_interrogatories': {
            'will_answer': 'Plaintiff will answer this interrogatory.',
            'cannot_answer': 'Plaintiff cannot answer this interrogatory at this time.',
            'no_text': ''
        }
    }
    
    responses_dict = _parse_ai_response_by_question(ai_response, type_config)
    combined_responses = RichText()
    
    # Get the appropriate response set based on discovery type
    is_special_interrogatory = type_config.get('display_name') == 'Special Interrogatories'
    response_set = standard_responses['special_interrogatories'] if is_special_interrogatory else standard_responses['requests_for_production']
    
    for question in questions:
        question_number = question.get('number', '')
        question_id = f"q_{question_number}"
        ai_response_text = responses_dict.get(question_number, 
            "No objections found. Subject to and without waiving the foregoing objections, Plaintiff responds as follows:")
        ai_response_text = strip_markdown(ai_response_text)
        user_selection = selections.get(question_id, 'no_text')
        standard_response = response_set.get(user_selection, "")
        
        # Format differently based on discovery type
        if is_special_interrogatory:
            # Special Interrogatories formatting
            combined_responses.add(f"\nSPECIAL INTERROGATORY NO. {question_number}:", bold=True, underline=True, font='Times New Roman')
            combined_responses.add(f"\n\t{question.get('text', '').strip()}", font='Times New Roman')
            combined_responses.add(f"\nRESPONSE TO SPECIAL INTERROGATORY NO. {question_number}:", bold=True, underline=True, font='Times New Roman')
        else:
            # Original RFP formatting (unchanged)
            combined_responses.add(f"\nREQUEST FOR PRODUCTION NO. {question_number}:", bold=True, underline=True, font='Times New Roman')
            combined_responses.add(f"\n\t{question.get('text', '').strip()}", font='Times New Roman')
            combined_responses.add(f"\nRESPONSE TO REQUEST FOR PRODUCTION NO. {question_number}:", bold=True, underline=True, font='Times New Roman')
        
        # Split objection and 'Subject to...' if present
        objection_part = ai_response_text
        subject_to_part = ""
        if "Subject to and without waiving" in ai_response_text:
            parts = ai_response_text.split("Subject to and without waiving", 1)
            objection_part = parts[0].strip()
            subject_to_part = "Subject to and without waiving" + parts[1]
        
        # Objection (first-line indent using tab)
        if objection_part:
            combined_responses.add(f"\n\t{objection_part}", font='Times New Roman')
        
        # 'Subject to and without waiving...' (first-line indent using tab)
        if subject_to_part:
            combined_responses.add(f"\n\t{subject_to_part}", font='Times New Roman')
        
        # Standard response (first-line indent using tab)
        if standard_response:
            combined_responses.add(f"\n\t{standard_response}", font='Times New Roman')
    
    return combined_responses

def _parse_ai_response_by_question(ai_response, type_config):
    """Parse AI response text to extract responses by question number."""
    responses_dict = {}
    current_num = None
    current_text = ""
    
    # Determine the response pattern based on discovery type
    response_pattern = type_config['response_type'].upper()
    
    for line in ai_response.split('\n'):
        if response_pattern in line:
            # Save previous response if exists
            if current_num is not None:
                responses_dict[current_num] = current_text.strip()
            
            # Extract question number
            try:
                current_num = line.split("NO.")[1].split(":")[0].strip()
                current_text = ""
            except (IndexError, ValueError):
                current_num = None
                
        elif type_config['request_type'].upper() in line:
            # If we hit a new request, save current response
            if current_num is not None:
                responses_dict[current_num] = current_text.strip()
                current_num = None
                
        elif current_num is not None:
            current_text += line + "\n"
    
    # Save last response
    if current_num is not None:
        responses_dict[current_num] = current_text.strip()
    
    print(f"DEBUG: Parsed {len(responses_dict)} AI responses")
    return responses_dict

def _render_and_send_document(template_name, context, case_id, discovery_type):
    """Render template and send document for download."""
    print(f"DEBUG: Rendering template: {template_name}")
    
    # Build template path
    template_dir = os.path.join(current_app.root_path, 'templates')
    template_path = os.path.join(template_dir, template_name)
    
    if not os.path.exists(template_path):
        print(f"ERROR: Template not found: {template_path}")
        return jsonify({'error': f'Template file "{template_name}" not found'}), 500
    
    try:
        # Create and render document
        doc = DocxTemplate(template_path)
        doc.render(context)
        
        # Save to memory
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        # Create filename
        safe_case_identifier = str(context.get('case_number', case_id)).replace('/', '_').replace('\\', '_')
        output_filename = f"{discovery_type}_responses_{safe_case_identifier}.docx"
        
        print(f"DEBUG: Sending file: {output_filename}")
        
        return send_file(
            output,
            as_attachment=True,
            download_name=output_filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        print(f"ERROR: Template rendering failed: {str(e)}")
        traceback.print_exc()
        return jsonify({'error': f'Failed to render template: {str(e)}'}), 500

@bp.route('/discovery/cases/<int:case_id>/respond', methods=['POST'])
@login_required
def respond_to_discovery(case_id):
    """
    Accepts a file upload and discovery_type, parses the file, processes with AI,
    and returns a Word document with formatted responses.
    """
    print(f"DEBUG: Discovery respond endpoint accessed by user {current_user.id} for case {case_id}")
    print(f"DEBUG: Request form data: {request.form}")
    print(f"DEBUG: Request files: {request.files.keys() if request.files else 'None'}")
    
    # Check for 'document' instead of 'file' to match Documents API pattern
    document = request.files.get('document')
    discovery_type = request.form.get('discovery_type')
    
    print(f"DEBUG: Extracted discovery_type: {discovery_type}")
    print(f"DEBUG: Extracted document: {document.filename if document else 'None'}")
    
    if not document or not discovery_type:
        print("DEBUG: Missing document or discovery_type, returning 400")
        return jsonify({'error': 'Missing document or discovery_type'}), 400

    # Create a temporary file with a secure name using tempfile module
    temp_file = None
    temp_path = None
    try:
        # Create a named temporary file that closes when done
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
            temp_path = temp_file.name
            print(f"DEBUG: Created temporary file at {temp_path}")
            
            # Save uploaded file content to the temporary file
            document.save(temp_file)
            temp_file.flush()  # Ensure all data is written
            
        print(f"DEBUG: File saved successfully, size: {os.path.getsize(temp_path) if os.path.exists(temp_path) else 'file not found'}")
        
        # Load case details
        print(f"DEBUG: Attempting to get case {case_id} for user {current_user.id}")
        case = get_case_by_id(case_id, user_id=current_user.id)
        print(f"DEBUG: Successfully retrieved case {case_id}")
        
        case_details = {}
        if hasattr(case, 'to_dict'):
            print("DEBUG: Case has to_dict method, calling it")
            case_details = case.to_dict()
        else:
            print("DEBUG: Case doesn't have to_dict method, creating basic dict")
            # Create a basic dict with essential attributes
            case_details = {
                'id': case.id,
                'display_name': case.display_name if hasattr(case, 'display_name') else 'Unknown',
                # Add other essential fields
            }
        
        # Load objection master sheet (as plain text)
        objection_master = get_cached_objection_master()
        print(f"DEBUG: Successfully loaded objection master sheet, length: {len(objection_master)}")

        # Use the orchestrator service
        print(f"DEBUG: About to initialize DiscoveryResponseService")
        service = DiscoveryResponseService()
        print(f"DEBUG: Service initialized successfully")
        
        print(f"DEBUG: Calling service.respond with discovery_type={discovery_type}")
        
        # Add some pre-processing checks to validate the PDF
        try:
            # Quick text extraction test to verify PDF is readable
            with fitz.open(temp_path) as doc:
                if len(doc) > 0:
                    sample_text = doc[0].get_text()[:200]
                    print(f"DEBUG: Sample text from first page: {sample_text}")
                else:
                    print("DEBUG: PDF appears to have no pages")
        except Exception as pdf_check_error:
            print(f"DEBUG: Error in PDF pre-check: {pdf_check_error}")
            
        # Get the AI response
        result = service.respond(discovery_type, temp_path, case_details, objection_master)
        print(f"DEBUG: Service.respond completed successfully")
        
        # Define path to the template
        template_name = 'discovery_responses_template.docx'  # Make sure this exists in templates folder
        template_dir = os.path.join(current_app.root_path, 'templates')
        template_path = os.path.join(template_dir, template_name)
        
        print(f"DEBUG: Looking for document template at: {template_path}")
        if not os.path.exists(template_path):
            return jsonify({'error': f'Template file "{template_name}" not found'}), 500
        
        # Extract questions and AI response from result
        questions = result.get('questions', [])
        ai_response = result.get('ai_response', '')
        
        # Parse the AI responses for relevant parts (assuming specific format)
        responses_dict = {}
        current_num = None
        current_text = ""
        in_response = False
        
        print(f"DEBUG: Parsing AI response text")
        
        # Parse AI response to match responses with question numbers
        for line in ai_response.split('\n'):
            if "RESPONSE TO REQUEST FOR PRODUCTION NO." in line:
                # If we were building a response, save it
                if current_num is not None:
                    responses_dict[current_num] = current_text.strip()
                
                # Extract the question number
                try:
                    current_num = line.split("NO.")[1].split(":")[0].strip()
                    current_text = ""
                except IndexError:
                    current_num = None
                    print(f"DEBUG: Failed to extract number from: {line}")
            elif "REQUEST FOR PRODUCTION NO." in line:
                # If we hit a new request, save the current response
                if current_num is not None:
                    responses_dict[current_num] = current_text.strip()
                    current_num = None
            elif current_num is not None:
                current_text += line + "\n"
        
        # Add the last response if we have one
        if current_num is not None:
            responses_dict[current_num] = current_text.strip()
        
        # Create RichText object with formatted responses
        print(f"DEBUG: Creating RichText for responses")
        responses_rt = RichText()
        
        for question in questions:
            question_number = question.get('number', '')
            
            # Get response text or use default, and strip markdown
            response_text = responses_dict.get(question_number, 
                          "No objections found. Subject to and without waiving the foregoing objections, Plaintiff responds as follows:")
            response_text = strip_markdown(response_text)
            
            # Get the user selection for this question
            question_id = f"q_{question_number}"
            selection = selections.get(question_id, 'no_text')
            standard_response = standard_responses.get(selection, "")
            
            # Add the AI response text
            responses_rt.add(response_text)
            
            # Add the selected standard response if it's not empty
            if standard_response:
                # If the response already ends with a period, add a space
                if response_text.strip().endswith('.'):
                    responses_rt.add(f" {standard_response}")
                else:
                    # Otherwise add a period then the standard response
                    responses_rt.add(f". {standard_response}")
            
            responses_rt.add("\n\n")
        
        # Create context with case info and responses
        context = {
            'case_name': case.display_name if hasattr(case, 'display_name') else '',
            'case_number': case.case_number if hasattr(case, 'case_number') else '',
            'plaintiff': case.plaintiff if hasattr(case, 'plaintiff') else '',
            'defendant': case.defendant if hasattr(case, 'defendant') else '',
            'judge': case.judge if hasattr(case, 'judge') else '',
            'jurisdiction': case.jurisdiction if hasattr(case, 'jurisdiction') else '',
            'county': case.county if hasattr(case, 'county') else '',
            # Date fields
            'filing_date': case.filing_date if hasattr(case, 'filing_date') else '',
            'trial_date': case.trial_date if hasattr(case, 'trial_date') else '',
            'incident_date': case.incident_date if hasattr(case, 'incident_date') else '',
            # Other fields
            'incident_location': case.incident_location if hasattr(case, 'incident_location') else '',
            'incident_description': case.incident_description if hasattr(case, 'incident_description') else '',
            'case_type': case.case_type if hasattr(case, 'case_type') else '',
            'defendant_counsel_info': case.defendant_counsel_info if hasattr(case, 'defendant_counsel_info') else '',
            'plaintiff_counsel_info': case.plaintiff_counsel_info if hasattr(case, 'plaintiff_counsel_info') else '',
            'vehicle_details': case.vehicle_details if hasattr(case, 'vehicle_details') else '',
            'defendant_counsel_attorneys': case.defendant_counsel_attorneys if hasattr(case, 'defendant_counsel_attorneys') else '',
            'defendant_counsel_firm': case.defendant_counsel_firm if hasattr(case, 'defendant_counsel_firm') else '',
            'defendant_counsel_address': case.defendant_counsel_address if hasattr(case, 'defendant_counsel_address') else '',
            'defendant_counsel_contact': case.defendant_counsel_contact if hasattr(case, 'defendant_counsel_contact') else '',
            'acting_attorney': case.acting_attorney if hasattr(case, 'acting_attorney') else '',
            'acting_clerk': case.acting_clerk if hasattr(case, 'acting_clerk') else '',
            # Add the responses last
            'responses': responses_rt
        }
        
        # Create and render the document
        print(f"DEBUG: Creating document from template")
        doc = DocxTemplate(template_path)
        doc.render(context)
        
        # Save to BytesIO
        print(f"DEBUG: Saving document to memory stream")
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        # Create filename
        safe_case_identifier = str(getattr(case, 'case_number', case_id)).replace('/','_').replace('\\','_')
        output_filename = f"RFP_Responses_{safe_case_identifier}.docx"
        
        # Delete the stored session data
        if session_key in current_app.config:
            del current_app.config[session_key]
        
        # Return the file
        print(f"DEBUG: Sending file: {output_filename}")
        return send_file(
            output,
            as_attachment=True,
            download_name=output_filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        error_trace = traceback.format_exc()
        print(f"DEBUG: Error in respond_to_discovery: {str(e)}")
        print(f"DEBUG: Error traceback: {error_trace}")
        current_app.logger.error(f"Error processing discovery for case {case_id}: {e}")
        current_app.logger.error(f"Traceback: {error_trace}")
        return jsonify({"error": f"An error occurred: {str(e)}"}), 500
    finally:
        # Clean up the temporary file
        if temp_file and temp_path and os.path.exists(temp_path):
            try:
                os.remove(temp_path)
                print(f"DEBUG: Temporary file {temp_path} removed")
            except Exception as cleanup_error:
                print(f"DEBUG: Error removing temporary file: {str(cleanup_error)}")

@bp.route('/discovery/interrogatory-questions', methods=['GET'])
@login_required
def get_interrogatory_questions():
    """
    Serves the list of interrogatory questions based on the requested language.
    """
    language = request.args.get('language', 'english')
    
    try:
        # Get the project root directory (one level up from the backend directory)
        project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..'))
        
        # Construct the path to the JSON file
        file_path = os.path.join(project_root, 'backend', 'data', f'form_interrogatories_{language}.json')
        
        print(f"DEBUG: Looking for questions file at: {file_path}")
        print(f"DEBUG: File exists? {os.path.exists(file_path)}")
        
        # Check if file exists
        if not os.path.exists(file_path):
            return jsonify({'error': f'No questions found for language: {language}'}), 404
            
        # Read and return the questions
        with open(file_path, 'r', encoding='utf-8') as f:
            questions = json.load(f)
            return jsonify(questions)
            
    except Exception as e:
        current_app.logger.error(f"Error serving interrogatory questions: {str(e)}")
        return jsonify({'error': 'Failed to load interrogatory questions'}), 500

@bp.route('/discovery/generate-interrogatory-document', methods=['POST'])
@login_required
def generate_interrogatory_document():
    """
    Generates an interrogatory document based on selected questions and case details.
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400

        # Extract required data
        case_id = data.get('case_id')
        selected_ids = data.get('selected_ids', [])
        language = data.get('language', 'english')

        if not case_id or not selected_ids:
            return jsonify({'error': 'Missing required fields: case_id and selected_ids'}), 400

        # Get case details
        case = get_case_by_id(case_id, user_id=current_user.id)
        if not case:
            return jsonify({'error': 'Case not found'}), 404

        # Get the questions from the JSON file
        project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..'))
        questions_file = os.path.join(project_root, 'backend', 'data', f'form_interrogatories_{language}.json')
        
        if not os.path.exists(questions_file):
            return jsonify({'error': f'No questions found for language: {language}'}), 404

        with open(questions_file, 'r', encoding='utf-8') as f:
            all_questions = json.load(f)
            
        # Filter questions based on selected IDs
        selected_questions = [q for q in all_questions if q['id'] in selected_ids]

        # Create a new Document
        doc = Document()
        
        # Add title
        doc.add_heading('Form Interrogatories', 0)
        
        # Add case information
        doc.add_heading('Case Information', level=1)
        doc.add_paragraph(f'Case: {case.display_name}')
        doc.add_paragraph(f'Case Number: {case.case_number if hasattr(case, "case_number") else "N/A"}')
        
        # Add selected questions
        doc.add_heading('Interrogatories', level=1)
        for question in selected_questions:
            # Add question number and text
            doc.add_heading(f'Interrogatory No. {question.get("number", "N/A")}', level=2)
            doc.add_paragraph(question.get('text', ''))
            
            # Add subparts if they exist
            if question.get('subparts'):
                for subpart in question['subparts']:
                    doc.add_paragraph(subpart, style='List Bullet')

        # Save to a temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()

        # Read the file and send it
        with open(temp_file.name, 'rb') as f:
            file_content = f.read()

        # Clean up the temporary file
        os.unlink(temp_file.name)

        # Create response with the file
        response = send_file(
            io.BytesIO(file_content),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'form_interrogatories_{language}_{case_id}.docx'
        )

        return response

    except Exception as e:
        error_trace = traceback.format_exc()
        print(f"DEBUG: Error generating interrogatory document: {str(e)}")
        print(f"DEBUG: Error traceback: {error_trace}")
        current_app.logger.error(f"Error generating interrogatory document: {e}")
        current_app.logger.error(f"Traceback: {error_trace}")
        return jsonify({'error': f'Failed to generate document: {str(e)}'}), 500

@bp.route('/discovery/cases/<int:case_id>/format-responses', methods=['POST'])
@login_required
def format_discovery_responses(case_id: int):
    """
    Format form interrogatory responses using AI.
    
    Args:
        case_id: ID of the case
        
    Returns:
        JSON response with formatted responses and metadata
    """
    try:
        # Get request data
        data = request.get_json()
        if not data or 'responses' not in data:
            return jsonify({'error': 'No responses provided'}), 400
            
        responses = data['responses']
        medical_data = data.get('medical_data', {})
        
        # Get case details
        case = Case.query.get_or_404(case_id)
        if case.user_id != current_user.id:
            return jsonify({'error': 'Unauthorized'}), 403
            
        # Convert case to dict for formatting
        case_dict = case_schema.dump(case)
        
        # Format responses using AI
        formatted_data = format_form_interrogatory_responses(responses, case_dict)
        
        # Format medical records if provided
        if medical_data:
            medical_formatted = format_medical_records(medical_data, case_dict)
            formatted_data['responses'].update(medical_formatted['responses'])
            formatted_data['metadata']['medical_records'] = medical_formatted['metadata']
        
        # Store in session for document generation
        session_key = f'formatted_responses_{case_id}'
        session[session_key] = {
            'responses': formatted_data['responses'],
            'formatted_at': datetime.utcnow().isoformat(),
            'case_id': case_id
        }
        
        return jsonify(formatted_data)
        
    except ValueError as e:
        logger.error(f"Validation error: {str(e)}")
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        logger.error(f"Error formatting responses: {str(e)}", exc_info=True)
        return jsonify({'error': 'Failed to format responses'}), 500